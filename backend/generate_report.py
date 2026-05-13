"""Convert docs/project-final-report.md to .docx and .pdf."""

import pathlib, re, sys

MD_PATH = pathlib.Path(__file__).parent.parent / "docs" / "project-final-report.md"
DOCX_PATH = MD_PATH.with_suffix(".docx")

# ── DOCX generation ──────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_table_from_md(doc, lines):
    """Parse a markdown table (list of raw lines) and add a Word table."""
    # Filter separator row
    rows = [l for l in lines if not re.match(r"^\|[-| :]+\|$", l.strip())]
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    cols = len(parsed[0])
    tbl = doc.add_table(rows=len(parsed), cols=cols)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = cell_text
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

text = MD_PATH.read_text(encoding="utf-8")
lines = text.splitlines()

i = 0
table_buf = []
in_table = False

while i < len(lines):
    line = lines[i]

    # Detect table start
    if line.strip().startswith("|"):
        if not in_table:
            in_table = True
            table_buf = []
        table_buf.append(line)
        i += 1
        continue
    else:
        if in_table:
            add_table_from_md(doc, table_buf)
            in_table = False
            table_buf = []

    stripped = line.strip()

    # Headings
    m = re.match(r"^(#{1,4})\s+(.*)", stripped)
    if m:
        level = len(m.group(1))
        add_heading(doc, m.group(2), level)
        i += 1
        continue

    # Horizontal rule
    if re.match(r"^---+$", stripped):
        doc.add_paragraph("─" * 60)
        i += 1
        continue

    # Image reference — add italicized caption
    img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
    if img_match:
        alt = img_match.group(1)
        p = doc.add_paragraph(f"[Figure: {alt}]")
        p.italic = True
        i += 1
        continue

    # Code block
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing ```
        p = doc.add_paragraph("\n".join(code_lines))
        p.style = "No Spacing"
        p.runs[0].font.name = "Courier New" if p.runs else "Courier New"
        continue

    # Inline code in normal paragraph (strip backticks)
    stripped_clean = re.sub(r"`([^`]+)`", r"\1", stripped)

    # Bold inline **text**
    # We emit plain text for simplicity
    stripped_clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", stripped_clean)

    # Bullet
    if stripped_clean.startswith("- "):
        doc.add_paragraph(stripped_clean[2:], style="List Bullet")
        i += 1
        continue

    # Empty line
    if not stripped_clean:
        doc.add_paragraph()
        i += 1
        continue

    # Normal paragraph
    doc.add_paragraph(stripped_clean)
    i += 1

# flush trailing table
if in_table:
    add_table_from_md(doc, table_buf)

doc.save(str(DOCX_PATH))
print(f"DOCX saved: {DOCX_PATH}")

# ── PDF via WeasyPrint or fallback ───────────────────────────────────────────
PDF_PATH = MD_PATH.with_suffix(".pdf")

try:
    import markdown as md_lib
    from weasyprint import HTML, CSS

    html_body = md_lib.markdown(text, extensions=["tables", "fenced_code"])
    html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; font-size: 11pt; margin: 2cm; }}
h1 {{ font-size: 18pt; }} h2 {{ font-size: 15pt; }} h3 {{ font-size: 13pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 12pt 0; }}
th, td {{ border: 1px solid #333; padding: 4pt 8pt; }}
th {{ background: #ddd; }}
pre {{ background: #f4f4f4; padding: 8pt; font-size: 9pt; }}
img {{ max-width: 100%; }}
</style>
</head><body>{html_body}</body></html>"""
    HTML(string=html_full).write_pdf(str(PDF_PATH))
    print(f"PDF saved:  {PDF_PATH}")
except ImportError:
    print("WeasyPrint not available — skipping PDF. Install with: pip install weasyprint")
except Exception as e:
    print(f"PDF generation failed: {e}")
